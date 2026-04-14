# doc_processor.py
"""
Word 文档占位符处理模块
支持解析 {{placeholder}} 格式的占位符，并用 RAG + LLM 生成内容填充
"""

import re
import os
from typing import Dict, List, Tuple
from docx import Document
from docx.shared import Pt
from utils.log_utils import log

# 占位符正则：匹配 {{xxx}} 格式
PLACEHOLDER_PATTERN = re.compile(r'\{\{([^}]+)\}\}')


def extract_placeholders(doc_path: str) -> List[str]:
    """
    从 Word 文档中提取所有占位符
    返回: 去重后的占位符名称列表
    """
    doc = Document(doc_path)
    placeholders = set()

    # 遍历所有段落
    for para in doc.paragraphs:
        matches = PLACEHOLDER_PATTERN.findall(para.text)
        placeholders.update(matches)

    # 遍历所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    matches = PLACEHOLDER_PATTERN.findall(para.text)
                    placeholders.update(matches)

    return list(placeholders)


def fill_placeholders(doc_path: str, replacements: Dict[str, str], output_path: str) -> Tuple[bool, str]:
    """
    用生成的内容填充 Word 文档中的占位符

    Args:
        doc_path: 原始文档路径
        replacements: {占位符名: 填充内容} 字典
        output_path: 输出文档路径

    Returns:
        (成功与否, 消息)
    """
    try:
        doc = Document(doc_path)
        filled_count = 0

        # 处理段落
        for para in doc.paragraphs:
            original_text = para.text
            new_text = original_text

            for placeholder, content in replacements.items():
                pattern = f'{{{{{placeholder}}}}}'  # {{xxx}}
                if pattern in new_text:
                    new_text = new_text.replace(pattern, content)
                    filled_count += 1

            # 如果有变化，更新段落（保持格式）
            if new_text != original_text:
                _replace_paragraph_text(para, new_text)

        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text
                        new_text = original_text

                        for placeholder, content in replacements.items():
                            pattern = f'{{{{{placeholder}}}}}'
                            if pattern in new_text:
                                new_text = new_text.replace(pattern, content)
                                filled_count += 1

                        if new_text != original_text:
                            _replace_paragraph_text(para, new_text)

        # 保存
        doc.save(output_path)
        log.info(f"文档填充完成，共填充 {filled_count} 处")
        return True, f"成功填充 {filled_count} 处内容"

    except Exception as e:
        log.error(f"文档填充失败: {e}")
        return False, str(e)


def _replace_paragraph_text(para, new_text: str):
    """
    替换段落文本，尽量保持原有格式
    """
    # 如果段落有 runs，保留第一个 run 的格式
    if para.runs:
        # 获取第一个 run 的格式
        first_run = para.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        bold = first_run.font.bold
        italic = first_run.font.italic

        # 清空所有 runs
        for run in para.runs:
            run.text = ""

        # 用第一个 run 设置新文本
        para.runs[0].text = new_text
        para.runs[0].font.name = font_name
        if font_size:
            para.runs[0].font.size = font_size
        para.runs[0].font.bold = bold
        para.runs[0].font.italic = italic
    else:
        para.text = new_text


def get_placeholder_context(doc_path: str, placeholder: str, context_chars: int = 200) -> str:
    """
    获取占位符周围的上下文，帮助 LLM 理解要填充什么内容
    """
    doc = Document(doc_path)

    for para in doc.paragraphs:
        if f'{{{{{placeholder}}}}}' in para.text:
            # 返回该段落的文本作为上下文
            return para.text

    # 检查表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if f'{{{{{placeholder}}}}}' in para.text:
                        return para.text

    return ""


def build_fill_prompt(placeholder: str, context: str, user_instruction: str = "") -> str:
    """
    构建用于生成填充内容的 prompt
    """
    prompt = f"""你需要为一个 Word 文档填充内容。

占位符名称: {placeholder}
占位符所在上下文: {context}
"""

    if user_instruction:
        prompt += f"\n用户要求: {user_instruction}\n"

    prompt += """
请根据占位符名称和上下文，生成合适的填充内容。
要求：
1. 内容要符合上下文语境
2. 语言风格要与文档一致
3. 只输出填充内容本身，不要加任何解释或前缀

填充内容："""

    return prompt